#!/usr/bin/env python3
"""
Examen Blanc Docx Parser
Parses docx files containing examen blanc questions and extracts structured data
"""

import os
import sys
import re
import json
from typing import List, Dict, Any, Optional
from docx import Document
from datetime import datetime

class ExamenBlancParser:
    def __init__(self):
        self.questions_parsed = 0
        self.errors = []
        
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Remove special characters that might cause issues
        text = text.replace('\u00a0', ' ')  # Non-breaking space
        text = text.replace('\u2013', '-')  # En dash
        text = text.replace('\u2014', '-')  # Em dash
        
        return text
    
    def extract_question_number(self, text: str) -> Optional[int]:
        """Extract question number from text"""
        # Look for patterns like "1.", "1)", "1-", etc.
        patterns = [
            r'^(\d+)\.',  # 1.
            r'^(\d+)\)',  # 1)
            r'^(\d+)-',   # 1-
            r'^(\d+)\s',  # 1 
            r'^(\d+)\t',  # 1\t (tab)
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text.strip())
            if match:
                return int(match.group(1))
        
        return None
    
    def is_option_line(self, text: str) -> bool:
        """Check if line is an option (a., b., c.)"""
        text = text.strip().lower()
        # Handle different option formats: a., a), A., A), etc.
        return bool(re.match(r'^[abc]\.?\s', text) or re.match(r'^[abc]\)\s', text))
    
    def extract_option(self, text: str) -> Optional[str]:
        """Extract option text from option line"""
        # Remove option letter and following punctuation (handle both . and ) formats)
        text = re.sub(r'^[abc]\.?\s*', '', text.strip())
        text = re.sub(r'^[abc]\)\s*', '', text.strip())
        return self.clean_text(text) if text else None
    
    def determine_category(self, question_number: int) -> str:
        """Determine category based on question number"""
        if 1 <= question_number <= 20:
            return "ANG"
        elif 21 <= question_number <= 40:
            return "CG"
        elif 41 <= question_number <= 60:
            return "LOG"
        else:
            return "UNKNOWN"
    
    def parse_docx_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse a single docx file and extract all questions"""
        questions = []
        
        try:
            doc = Document(file_path)
            
            print(f"🔍 Parsing {os.path.basename(file_path)} - {len(doc.paragraphs)} paragraphs")
            
            current_section = None
            question_number = 1
            i = 0
            
            while i < len(doc.paragraphs):
                paragraph = doc.paragraphs[i]
                text = self.clean_text(paragraph.text)
                
                if not text:
                    i += 1
                    continue
                
                # Debug: Show first few non-empty paragraphs
                if i < 10:
                    print(f"  Para {i}: {text[:100]}...")
                
                # Check for section headers
                if text.startswith("SECTION") or text.startswith("TEST") or text.startswith("🎓"):
                    if "CULTURE" in text.upper() or "GÉNÉRALE" in text.upper():
                        current_section = "CG"
                        print(f"  📚 Found CG section: {text}")
                    elif "ANGLAIS" in text.upper():
                        current_section = "ANG"
                        print(f"  🇬🇧 Found ANG section: {text}")
                    elif "LOGIQUE" in text.upper() or "NUMÉRIQUE" in text.upper():
                        current_section = "LOG"
                        print(f"  🧮 Found LOG section: {text}")
                    i += 1
                    continue
                
                # Check if this looks like a question with options (single paragraph)
                if self.looks_like_question_with_options(text):
                    # Use fallback category if no section detected
                    section = current_section if current_section else self.determine_category(question_number)
                    print(f"  Found single-paragraph question {question_number} in {section}: {text[:50]}...")
                    
                    # Parse the question and options
                    parsed = self.parse_question_with_options(text, question_number, file_path, section)
                    if parsed:
                        questions.append(parsed)
                        self.questions_parsed += 1
                        print(f"  ✅ Saved question {question_number} as {section}")
                        question_number += 1
                    i += 1
                    continue
                
                # Check if this looks like a question number (e.g., "1. Question text :")
                if self.looks_like_question_number(text) and i + 3 < len(doc.paragraphs):
                    # Check if next 3 paragraphs are options
                    option_paragraphs = []
                    for j in range(1, 4):  # Check next 3 paragraphs
                        if i + j < len(doc.paragraphs):
                            next_text = self.clean_text(doc.paragraphs[i + j].text)
                            if self.looks_like_single_option(next_text):
                                option_paragraphs.append(next_text)
                    
                    # If we found 3 options, this is a multi-line question
                    if len(option_paragraphs) == 3:
                        # Combine question and options
                        combined_text = text + " " + " ".join(option_paragraphs)
                        
                        # Use fallback category if no section detected
                        section = current_section if current_section else self.determine_category(question_number)
                        print(f"  Found multi-line question {question_number} in {section}: {text[:50]}...")
                        
                        # Parse the combined question and options
                        parsed = self.parse_question_with_options(combined_text, question_number, file_path, section)
                        if parsed:
                            questions.append(parsed)
                            self.questions_parsed += 1
                            print(f"  ✅ Saved question {question_number} as {section}")
                            question_number += 1
                        i += 4  # Skip question + 3 options
                        continue
                
                # Check if this looks like a question without options (multi-paragraph format)
                if self.looks_like_question_text(text) and i + 1 < len(doc.paragraphs):
                    next_paragraph = doc.paragraphs[i + 1]
                    next_text = self.clean_text(next_paragraph.text)
                    
                    # Check if next paragraph has options
                    if self.looks_like_options_only(next_text):
                        # Combine question and options
                        combined_text = text + " " + next_text
                        
                        # Use fallback category if no section detected
                        section = current_section if current_section else self.determine_category(question_number)
                        print(f"  Found multi-paragraph question {question_number} in {section}: {text[:50]}...")
                        
                        # Parse the combined question and options
                        parsed = self.parse_question_with_options(combined_text, question_number, file_path, section)
                        if parsed:
                            questions.append(parsed)
                            self.questions_parsed += 1
                            print(f"  ✅ Saved question {question_number} as {section}")
                            question_number += 1
                        i += 2  # Skip both paragraphs
                        continue
                
                i += 1
            
            print(f"✅ Parsed {len(questions)} questions from {os.path.basename(file_path)}")
            
        except Exception as e:
            error_msg = f"Error parsing {file_path}: {str(e)}"
            self.errors.append(error_msg)
            print(f"❌ {error_msg}")
        
        return questions
    
    def looks_like_question_with_options(self, text: str) -> bool:
        """Check if text looks like a question with options"""
        # Check for both formats: a. b. c. and a) b) c)
        return (bool(re.search(r'\ba\.\s+.*\bb\.\s+.*\bc\.\s+', text, re.IGNORECASE | re.DOTALL)) or
                bool(re.search(r'\ba\)\s+.*\bb\)\s+.*\bc\)\s+', text, re.IGNORECASE | re.DOTALL)))
    
    def looks_like_question_text(self, text: str) -> bool:
        """Check if text looks like a question without options"""
        # Look for question patterns: ends with ? or : and has reasonable length
        # Also check if it doesn't contain options (a. b. c.)
        has_options = bool(re.search(r'\ba\.\s+.*\bb\.\s+.*\bc\.\s+', text, re.IGNORECASE | re.DOTALL)) or bool(re.search(r'\ba\)\s+.*\bb\)\s+.*\bc\)\s+', text, re.IGNORECASE | re.DOTALL))
        
        return (text.endswith('?') or text.endswith(':')) and len(text) > 20 and not text.startswith("SECTION") and not text.startswith("TEST") and not text.startswith("🎓") and not has_options
    
    def looks_like_question_number(self, text: str) -> bool:
        """Check if text looks like a question number (e.g., '1. Question text :')"""
        # Look for pattern: number. question text ending with : or ?
        return bool(re.match(r'^\d+\.\s+.*[?:]\s*$', text.strip()))
    
    def looks_like_single_option(self, text: str) -> bool:
        """Check if text looks like a single option (e.g., 'a) option text' or 'a. option text')"""
        # Look for pattern that starts with a), b), c) or a., b., c.
        return bool(re.match(r'^[abc][\)\.]\s+', text.strip(), re.IGNORECASE))
    
    def looks_like_options_only(self, text: str) -> bool:
        """Check if text looks like options only (a. b. c. pattern)"""
        # Look for pattern that starts with a. and has b. and c.
        return bool(re.search(r'^a\.\s+.*\bb\.\s+.*\bc\.\s+', text, re.IGNORECASE | re.DOTALL)) or bool(re.search(r'^a\)\s+.*\bb\)\s+.*\bc\)\s+', text, re.IGNORECASE | re.DOTALL))
    
    def parse_question_with_options(self, text: str, question_number: int, file_path: str, section: str) -> Optional[Dict[str, Any]]:
        """Parse a single line containing question and options"""
        try:
            # Handle questions with line breaks by normalizing the text first
            # Replace line breaks with spaces, but preserve the structure
            normalized_text = re.sub(r'\s+', ' ', text.strip())
            
            # More robust approach: find the actual option pattern
            # Look for the pattern: question text a. option1 b. option2 c. option3
            # Use word boundaries to ensure we find actual option markers, not letters in question text
            
            # Find the first occurrence of "a. " or "a) " that's followed by text, then "b. " or "b) ", then "c. " or "c) "
            a_match = re.search(r'\ba[\.\)]\s+', normalized_text, re.IGNORECASE)
            if not a_match:
                return None
            
            a_pos = a_match.start()
            a_marker_len = len(a_match.group())
            
            # Find "b. " or "b) " that comes after "a. " or "a) "
            b_match = re.search(r'\bb[\.\)]\s+', normalized_text[a_pos:], re.IGNORECASE)
            if not b_match:
                return None
            
            b_pos = a_pos + b_match.start()
            b_marker_len = len(b_match.group())
            
            # Find "c. " or "c) " that comes after "b. " or "b) "
            c_match = re.search(r'\bc[\.\)]\s+', normalized_text[b_pos:], re.IGNORECASE)
            if not c_match:
                return None
            
            c_pos = b_pos + c_match.start()
            c_marker_len = len(c_match.group())
            
            # Ensure they're in order
            if not (a_pos < b_pos < c_pos):
                return None
            
            # Extract question text (everything before "a." or "a)")
            question_text = normalized_text[:a_pos].strip()
            
            # Extract options
            answer1 = normalized_text[a_pos+a_marker_len:b_pos].strip()
            answer2 = normalized_text[b_pos+b_marker_len:c_pos].strip()
            answer3 = normalized_text[c_pos+c_marker_len:].strip()
            
            # Clean up the options (remove any trailing punctuation that might be part of the next question)
            answer1 = answer1.rstrip('.,;:')
            answer2 = answer2.rstrip('.,;:')
            answer3 = answer3.rstrip('.,;:')
            
            return {
                "question_text": question_text,
                "answer1": answer1,
                "answer2": answer2,
                "answer3": answer3,
                "correct": None,  # Will be filled by AI
                "explanation": "",  # Will be filled by AI
                "category": section or "UNKNOWN",
                "difficulty": "HARD",
                "exam_type": "CMS" if "cms" in file_path.lower() else "CS",
                "test_type": "examen_blanc",
                "sub_category": None,
                "unique_hash": f"examen_blanc_{question_number}_{hash(text) % 1000000}",
                "ai_generated": True,
                "question_pool": f"{section}_examen_blanc",
                "usage_count": 0,
                "is3Option": True,
                "file_source": os.path.basename(file_path),
                "question_number": question_number
            }
            
        except Exception as e:
            error_msg = f"Error parsing question {question_number}: {str(e)}"
            self.errors.append(error_msg)
            return None
    
    def extract_options_from_text(self, options_text: str) -> List[str]:
        """Extract individual options from text like 'a. option1 b. option2 c. option3'"""
        options = []
        
        # Try different patterns - more flexible approach
        patterns = [
            r'[abc]\.\s+([^abc]+?)(?=\s+[abc]\.|$)',  # a. option1 b. option2 c. option3
            r'[abc]\)\s+([^abc]+?)(?=\s+[abc]\)|$)',  # a) option1 b) option2 c) option3
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, options_text, re.IGNORECASE)
            if len(matches) == 3:
                return [match.strip() for match in matches]
        
        # Fallback: manual split approach
        text = options_text.strip()
        
        # Handle a. format
        if text.startswith("a. "):
            text = text[3:]  # Remove "a. "
            parts = text.split(" b. ")
            if len(parts) == 2:
                option1 = parts[0]
                remaining = parts[1]
                parts2 = remaining.split(" c. ")
                if len(parts2) == 2:
                    option2 = parts2[0]
                    option3 = parts2[1]
                    return [option1.strip(), option2.strip(), option3.strip()]
        
        # Handle a) format
        elif text.startswith("a) "):
            text = text[3:]  # Remove "a) "
            parts = text.split(" b) ")
            if len(parts) == 2:
                option1 = parts[0]
                remaining = parts[1]
                parts2 = remaining.split(" c) ")
                if len(parts2) == 2:
                    option2 = parts2[0]
                    option3 = parts2[1]
                    return [option1.strip(), option2.strip(), option3.strip()]
        
        return []
    
    def parse_all_files(self, base_path: str) -> Dict[str, List[Dict[str, Any]]]:
        """Parse all docx files in the new_exams directory"""
        all_questions = {}
        
        # Process CMS exams
        cms_path = os.path.join(base_path, "cms_exams")
        if os.path.exists(cms_path):
            print(f"\n📚 Processing CMS exams from: {cms_path}")
            cms_files = [f for f in os.listdir(cms_path) if f.endswith('.docx')]
            cms_files.sort()  # Sort for consistent processing order
            
            for filename in cms_files:
                file_path = os.path.join(cms_path, filename)
                questions = self.parse_docx_file(file_path)
                all_questions[filename] = questions
        
        # Process CS exams
        cs_path = os.path.join(base_path, "cs_exams")
        if os.path.exists(cs_path):
            print(f"\n📚 Processing CS exams from: {cs_path}")
            cs_files = [f for f in os.listdir(cs_path) if f.endswith('.docx')]
            cs_files.sort()  # Sort for consistent processing order
            
            for filename in cs_files:
                file_path = os.path.join(cs_path, filename)
                questions = self.parse_docx_file(file_path)
                all_questions[filename] = questions
        
        return all_questions
    
    def save_parsed_data(self, all_questions: Dict[str, List[Dict[str, Any]]], output_dir: str):
        """Save parsed questions to JSON files"""
        os.makedirs(output_dir, exist_ok=True)
        
        for filename, questions in all_questions.items():
            # Create output filename
            base_name = os.path.splitext(filename)[0]
            output_file = os.path.join(output_dir, f"{base_name}_parsed.json")
            
            # Prepare data for JSON
            output_data = {
                'filename': filename,
                'parsed_at': datetime.now().isoformat(),
                'total_questions': len(questions),
                'questions': questions
            }
            
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(output_data, f, indent=2, ensure_ascii=False)
                print(f"💾 Saved {len(questions)} questions to {output_file}")
            except Exception as e:
                error_msg = f"Error saving {output_file}: {str(e)}"
                self.errors.append(error_msg)
                print(f"❌ {error_msg}")
    
    def generate_summary_report(self, all_questions: Dict[str, List[Dict[str, Any]]]) -> str:
        """Generate a summary report of parsed data"""
        total_files = len(all_questions)
        total_questions = sum(len(questions) for questions in all_questions.values())
        
        # Count by category
        category_counts = {"ANG": 0, "CG": 0, "LOG": 0, "UNKNOWN": 0}
        for questions in all_questions.values():
            for q in questions:
                category = q.get('category', 'UNKNOWN')
                if category is None:
                    category = 'UNKNOWN'
                category_counts[category] += 1
        
        # Count by exam type
        exam_type_counts = {"CMS": 0, "CS": 0}
        for filename in all_questions.keys():
            if "cms" in filename.lower():
                exam_type_counts["CMS"] += len(all_questions[filename])
            elif "cs" in filename.lower():
                exam_type_counts["CS"] += len(all_questions[filename])
        
        report = f"""
PARSING SUMMARY REPORT
======================
Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

Files Processed: {total_files}
Total Questions: {total_questions}

By Exam Type:
- CMS: {exam_type_counts['CMS']} questions
- CS: {exam_type_counts['CS']} questions

By Category:
- ANG: {category_counts['ANG']} questions
- CG: {category_counts['CG']} questions  
- LOG: {category_counts['LOG']} questions
- UNKNOWN: {category_counts['UNKNOWN']} questions

Expected vs Actual:
- Expected: {total_files * 60} questions (40 files × 60 questions each)
- Actual: {total_questions} questions
- Difference: {total_questions - (total_files * 60)} questions

Errors: {len(self.errors)}
"""
        
        if self.errors:
            report += "\nErrors:\n"
            for error in self.errors:
                report += f"- {error}\n"
        
        return report

def main():
    """Main function to run the parser"""
    if len(sys.argv) != 2:
        print("Usage: python parse_examen_blanc_docx.py <new_exams_directory>")
        print("Example: python parse_examen_blanc_docx.py /Users/joasyepidan/Documents/projects/ena/new_exams")
        sys.exit(1)
    
    new_exams_path = sys.argv[1]
    
    if not os.path.exists(new_exams_path):
        print(f"❌ Directory not found: {new_exams_path}")
        sys.exit(1)
    
    print("🔍 Starting examen blanc docx parsing...")
    print(f"📁 Processing directory: {new_exams_path}")
    
    parser = ExamenBlancParser()
    
    # Parse all files
    all_questions = parser.parse_all_files(new_exams_path)
    
    # Save parsed data
    output_dir = "/Users/joasyepidan/Documents/projects/ena/project/parsed_questions"
    parser.save_parsed_data(all_questions, output_dir)
    
    # Generate and save summary report
    report = parser.generate_summary_report(all_questions)
    print(report)
    
    # Save report to file
    report_file = "/Users/joasyepidan/Documents/projects/ena/project/parsing_summary_report.txt"
    try:
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(report)
        print(f"📄 Summary report saved to: {report_file}")
    except Exception as e:
        print(f"❌ Error saving report: {e}")
    
    print(f"\n🎉 Parsing complete!")
    print(f"✅ Total questions parsed: {parser.questions_parsed}")
    print(f"❌ Errors encountered: {len(parser.errors)}")

if __name__ == "__main__":
    main()
